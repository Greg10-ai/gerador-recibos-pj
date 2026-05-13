import os
import io
import pandas as pd
import unicodedata
import zipfile

from flask import Flask, render_template, request, send_file, redirect, url_for, session
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from num2words import num2words

# ==========================
# GERAÇÃO DE IMAGEM — SEM PLAYWRIGHT
# Usa WeasyPrint (HTML→PDF) + pdf2image (PDF→PNG)
# ou fallback para matplotlib se WeasyPrint falhar
# ==========================
def _render_html_to_png_weasyprint(html: str, output_path: str) -> bool:
    """Tenta renderizar HTML como PNG via WeasyPrint + pdf2image."""
    try:
        from weasyprint import HTML
        from pdf2image import convert_from_bytes

        pdf_bytes = HTML(string=html).write_pdf()
        pages = convert_from_bytes(pdf_bytes, dpi=150)
        if pages:
            pages[0].save(output_path, "PNG")
            return True
    except Exception as e:
        print(f"⚠️  WeasyPrint/pdf2image falhou: {e}")
    return False


def _render_table_to_png_matplotlib(
    descricao_col: str,
    valor_col: str,
    rows: list,          # lista de (desc, valor_str, is_total)
    mes_escolhido: str,
    output_path: str,
) -> bool:
    """Fallback: gera a tabela como imagem usando apenas matplotlib."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        from matplotlib.patches import FancyBboxPatch

        AZUL_ESCURO  = "#1F3864"
        AZUL_CLARO   = "#D9E1F2"
        BRANCO       = "#FFFFFF"
        TEXTO_ESCURO = "#1a1a1a"

        n = len(rows)
        altura_linha = 0.45
        margem_topo  = 1.2
        fig_h        = margem_topo + n * altura_linha + 0.5
        fig_w        = 7

        fig, ax = plt.subplots(figsize=(fig_w, fig_h))
        ax.axis("off")
        fig.patch.set_facecolor(BRANCO)

        # Título
        ax.text(
            0.5, 1.0 - (0.3 / fig_h),
            f"Apuração Mês {mes_escolhido.capitalize()}/{datetime.now().year}",
            transform=ax.transAxes,
            fontsize=13, fontweight="bold", color=AZUL_ESCURO,
            ha="center", va="top",
        )

        col_w   = [0.72, 0.28]   # proporções de largura das colunas
        x0      = 0.02
        y_start = 1.0 - (margem_topo / fig_h)

        def draw_cell(ax, x, y, w, h, text, bg, fg, bold=False, align="left"):
            rect = FancyBboxPatch(
                (x, y - h), w, h,
                boxstyle="square,pad=0",
                linewidth=0.5, edgecolor="#b0b8cc",
                facecolor=bg, transform=ax.transAxes, clip_on=False,
            )
            ax.add_patch(rect)
            tx = x + 0.01 if align == "left" else x + w - 0.01
            ha = "left"    if align == "left" else "right"
            ax.text(
                tx, y - h / 2, text,
                transform=ax.transAxes,
                fontsize=8.5, color=fg,
                fontweight="bold" if bold else "normal",
                ha=ha, va="center", clip_on=False,
            )

        h_row = altura_linha / fig_h

        # Cabeçalho
        draw_cell(ax, x0,              y_start, col_w[0], h_row, descricao_col, AZUL_ESCURO, BRANCO, bold=True, align="left")
        draw_cell(ax, x0 + col_w[0],  y_start, col_w[1], h_row, str(valor_col), AZUL_ESCURO, BRANCO, bold=True, align="right")

        for idx, (desc, val, is_total) in enumerate(rows):
            y = y_start - (idx + 1) * h_row
            if is_total:
                bg, fg, bold = AZUL_ESCURO, BRANCO, True
            elif idx % 2 == 0:
                bg, fg, bold = BRANCO, TEXTO_ESCURO, False
            else:
                bg, fg, bold = AZUL_CLARO, TEXTO_ESCURO, False

            draw_cell(ax, x0,             y, col_w[0], h_row, str(desc), bg, fg, bold=bold, align="left")
            draw_cell(ax, x0 + col_w[0], y, col_w[1], h_row, str(val),  bg, fg, bold=bold, align="right")

        plt.tight_layout(pad=0)
        plt.savefig(output_path, dpi=150, bbox_inches="tight", facecolor=BRANCO)
        plt.close(fig)
        return True
    except Exception as e:
        print(f"❌ matplotlib fallback falhou: {e}")
    return False


app = Flask(__name__)
app.secret_key = "segredo123"

UPLOAD_FOLDER = "/tmp/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

USUARIO = "Evelyn"
SENHA   = "Monique"

# ==========================
# NORMALIZAR NOMES
# ==========================
def normalizar_nome(nome):
    nome = str(nome).strip().lower()
    nome = unicodedata.normalize('NFKD', nome).encode('ASCII', 'ignore').decode('ASCII')
    nome = " ".join(nome.split())
    nome = nome.replace(" ", "")
    return nome

# ==========================
# DADOS BANCÁRIOS
# ==========================
def carregar_dados_bancarios(caminho):
    df = pd.read_excel(caminho)
    df.columns = [normalizar_nome(col) for col in df.columns]

    mapa = {}
    for _, row in df.iterrows():
        nome = normalizar_nome(row.get("vendedores", ""))
        info = {
            "nome":    row.get("vendedores", ""),
            "cnpj":    row.get("cnpj", ""),
            "banco":   row.get("banco", ""),
            "agencia": row.get("agencia", ""),
            "conta":   row.get("conta", ""),
            "pix":     row.get("pix", "")
        }
        mapa[nome] = info

    return mapa

def encontrar_dados_bancarios(nome_vendedor, mapa_banco):
    chave = normalizar_nome(nome_vendedor)
    if chave in mapa_banco:
        return mapa_banco[chave]
    for k in mapa_banco:
        if chave in k or k in chave:
            print(f"⚠️ Banco match aproximado: {nome_vendedor} -> {k}")
            return mapa_banco[k]
    print(f"❌ Sem dados bancários para: {nome_vendedor}")
    return {}

def encontrar_imagem(nome_vendedor, mapa_imagens):
    chave = normalizar_nome(nome_vendedor)
    if chave in mapa_imagens:
        return mapa_imagens[chave]
    for k in mapa_imagens:
        if chave in k or k in chave:
            return mapa_imagens[k]
    return None

# ==========================
# UTILITÁRIOS
# ==========================
def tratar_valor(valor):
    if pd.isna(valor):
        return None
    if isinstance(valor, str):
        valor = valor.replace("%", "").replace(".", "").replace(",", ".").strip()
    try:
        return round(float(valor), 2)
    except:
        return None

def formatar_real(valor):
    valor = round(float(valor), 2)
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def formatar_data_extenso(data_str):
    meses = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    data = datetime.strptime(data_str, "%Y-%m-%d")
    return f"Porto Alegre, {data.day} de {meses[data.month - 1]} de {data.year}"

def competencia_mes(mes):
    return f"{mes}/{datetime.now().year}"

# ==========================
# GERAR IMAGENS DAS ABAS
# Substitui Playwright por WeasyPrint + pdf2image (com fallback para matplotlib)
# ==========================
CAMPOS_PERCENTUAL = {
    "icm meta", "icm novos", "meta margem", "real",
    "icm", "icm meta base ativa", "% carteira ativa", "% liquidado"
}

def formatar_celula(v, descricao=""):
    desc_lower = str(descricao).strip().lower()
    eh_percentual = desc_lower in CAMPOS_PERCENTUAL

    try:
        f = float(v)
        if eh_percentual:
            return f"{f * 100:.2f}%" if abs(f) <= 1.5 else f"{f:.2f}%"
        if f == int(f):
            return f"{int(f):,}".replace(",", ".")
        return f"R$ {f:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return str(v) if pd.notna(v) else "-"


def gerar_imagens_abas(caminho_excel, mes_escolhido):
    imagens = {}
    xls = pd.ExcelFile(caminho_excel)

    for aba in xls.sheet_names:
        try:
            df_full = pd.read_excel(xls, sheet_name=aba, header=None)

            mes_lower = mes_escolhido.lower()
            linha_mes = col_mes = None

            for i in range(len(df_full)):
                linha = df_full.iloc[i].fillna("").astype(str).str.lower()
                if any(mes_lower in cel for cel in linha):
                    linha_mes = i
                    for j, cel in enumerate(linha):
                        if mes_lower in cel:
                            col_mes = j
                            break
                    break

            if linha_mes is None or col_mes is None:
                print(f"⚠️ Mês '{mes_escolhido}' não encontrado na aba '{aba}'")
                continue

            header_row = linha_mes + 1
            df = pd.read_excel(xls, sheet_name=aba, header=header_row)

            descricao_col = df.columns[1]
            valor_col     = df.columns[col_mes]

            df_ab = df[[descricao_col, valor_col]].dropna(how="all").head(60)

            # Prepara linhas formatadas
            rows_fmt = []
            linhas_html = ""
            for _, row in df_ab.iterrows():
                desc    = row[descricao_col]
                val_str = formatar_celula(row[valor_col], desc)
                is_tot  = "total" in str(desc).lower()
                rows_fmt.append((desc, val_str, is_tot))

                estilo_tr   = 'class="total"' if is_tot else ""
                linhas_html += f"<tr {estilo_tr}><td>{desc}</td><td>{val_str}</td></tr>\n"

            caminho_img = os.path.join(UPLOAD_FOLDER, f"{aba}.png")

            # ── Tenta WeasyPrint primeiro ──────────────────────────────────
            html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
        font-family: Arial, sans-serif;
        background: #ffffff;
        padding: 24px;
        width: 480px;
    }}
    h3 {{ color: #1F3864; font-size: 14px; margin-bottom: 12px; font-weight: bold; }}
    table {{ border-collapse: collapse; width: 100%; font-size: 12px; }}
    thead tr th {{
        background-color: #1F3864; color: #ffffff;
        padding: 7px 12px; text-align: center;
        font-weight: bold; border: 1px solid #1F3864;
    }}
    tbody tr td {{ padding: 5px 12px; border: 1px solid #d0d7e8; color: #1a1a1a; }}
    tbody tr td:first-child {{ text-align: left; font-weight: 500; }}
    tbody tr td:last-child  {{ text-align: right; }}
    tbody tr:nth-child(even) {{ background-color: #D9E1F2; }}
    tbody tr:nth-child(odd)  {{ background-color: #ffffff; }}
    tbody tr.total td {{
        background-color: #1F3864 !important;
        color: #ffffff !important;
        font-weight: bold; font-size: 13px;
    }}
</style>
</head>
<body>
    <h3>Apuração Mês {mes_escolhido.capitalize()}/{datetime.now().year}</h3>
    <table>
        <thead>
            <tr><th>{descricao_col}</th><th>{valor_col}</th></tr>
        </thead>
        <tbody>{linhas_html}</tbody>
    </table>
</body>
</html>"""

            ok = _render_html_to_png_weasyprint(html, caminho_img)

            # ── Fallback: matplotlib (sem dependências de sistema) ─────────
            if not ok:
                print(f"↩️  Usando matplotlib para '{aba}'")
                ok = _render_table_to_png_matplotlib(
                    str(descricao_col), str(valor_col),
                    rows_fmt, mes_escolhido, caminho_img
                )

            if ok:
                imagens[normalizar_nome(aba)] = caminho_img
                print(f"✅ Imagem gerada: {aba} -> {caminho_img}")
            else:
                print(f"❌ Não foi possível gerar imagem para '{aba}'")

        except Exception as e:
            print(f"❌ Erro na aba '{aba}': {e}")
            continue

    return imagens

# ==========================
# GERAR RECIBO
# ==========================
def gerar_recibo(vendedor, dados, mes, total, data_recibo, imagem=None, dados_bancarios=None):
    doc = Document()

    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = formatar_data_extenso(data_recibo)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading("RECIBO DE PAGAMENTO", level=1)

    total_extenso = num2words(total, lang='pt_BR').capitalize()

    doc.add_paragraph(
        f"\nInformo que recebi da empresa 3i Importação e Exportação Ltda, CNPJ 20.783.843/0001-19, "
        f"o valor de {formatar_real(total)} "
        f"({total_extenso}), "
        f"referente a serviços prestados de gestão comercial."
    )

    doc.add_paragraph(f"\nCompetência: {competencia_mes(mes)}")

    doc.add_heading("PROVENTOS", level=2)
    for desc, valor in dados.items():
        if valor > 0:
            doc.add_paragraph(f"{desc}: {formatar_real(valor)}")

    doc.add_heading("DESCONTOS", level=2)
    for desc, valor in dados.items():
        if valor < 0:
            doc.add_paragraph(f"{desc}: {formatar_real(valor)}")

    doc.add_heading(f"TOTAL LÍQUIDO: {formatar_real(total)}", level=2)

    if dados_bancarios:
        doc.add_paragraph("")
        doc.add_heading("DADOS BANCÁRIOS", level=2)
        doc.add_paragraph(f"Nome: {dados_bancarios.get('nome', '')}")
        doc.add_paragraph(f"CNPJ: {dados_bancarios.get('cnpj', '')}")
        doc.add_paragraph(f"Banco: {dados_bancarios.get('banco', '')}")
        doc.add_paragraph(f"Agência: {dados_bancarios.get('agencia', '')}")
        doc.add_paragraph(f"Conta: {dados_bancarios.get('conta', '')}")
        doc.add_paragraph(f"PIX: {dados_bancarios.get('pix', '')}")

    if imagem and os.path.exists(imagem):
        # ✅ FIX 1: page break embutido no parágrafo do título, sem criar parágrafo vazio extra
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from PIL import Image as PILImage

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_break = titulo.add_run()
        # Insere a quebra de página dentro do próprio run, sem parágrafo em branco
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run_break._r.append(br)
        run_break.text = f"Apuração Mês {mes}/{datetime.now().year}"

        # ✅ FIX 2: calcula largura/altura para caber na página sem cortar
        # Página A4: 21cm úteis com margens de ~2.5cm = ~16cm = 6.3 inches de largura
        # Altura útil por página: ~25cm = ~9.8 inches
        MAX_W = 6.0   # inches
        MAX_H = 8.5   # inches — deixa espaço para o título acima

        with PILImage.open(imagem) as img:
            img_w_px, img_h_px = img.size

        # DPI da imagem gerada foi 150
        img_w_in = img_w_px / 150
        img_h_in = img_h_px / 150

        # Escala para caber dentro dos limites mantendo proporção
        escala = min(MAX_W / img_w_in, MAX_H / img_h_in, 1.0)
        final_w = Inches(img_w_in * escala)

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(imagem, width=final_w)

    caminho = os.path.join(UPLOAD_FOLDER, f"Recibo_{vendedor}.docx")
    doc.save(caminho)
    return caminho

# ==========================
# ROTAS
# ==========================
@app.route("/")
def home():
    return render_template("login.html")

@app.route("/login", methods=["POST"])
def login():
    if request.form["usuario"] == USUARIO and request.form["senha"] == SENHA:
        session["logado"] = True
        return redirect(url_for("sistema"))
    return "Login inválido"

@app.route("/sistema", methods=["GET", "POST"])
def sistema():
    if not session.get("logado"):
        return redirect(url_for("home"))

    meses = ["Janeiro","Fevereiro","Março","Abril","Maio","Junho",
             "Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"]
    mes_atual = meses[datetime.now().month - 1]

    if request.method == "POST":
        try:
            data_recibo = request.form.get("data_recibo")

            file = request.files["file"]
            caminho = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(caminho)

            mapa_banco = {}
            file_banco = request.files.get("file_banco")
            if file_banco and file_banco.filename != "":
                caminho_banco = os.path.join(UPLOAD_FOLDER, file_banco.filename)
                file_banco.save(caminho_banco)
                mapa_banco = carregar_dados_bancarios(caminho_banco)

            mapa_imagens = {}
            file_imagens = request.files.get("file_imagens")
            if file_imagens and file_imagens.filename != "":
                caminho_img_excel = os.path.join(UPLOAD_FOLDER, file_imagens.filename)
                file_imagens.save(caminho_img_excel)
                mapa_imagens = gerar_imagens_abas(caminho_img_excel, request.form["mes"])

            df_full = pd.read_excel(caminho, header=None)
            mes_escolhido = request.form["mes"].lower()

            linha_mes = col_inicio = col_fim = None

            for i in range(len(df_full)):
                linha = df_full.iloc[i].fillna("").astype(str).str.lower()
                if any(mes_escolhido in cel for cel in linha):
                    linha_mes = i
                    for j, cel in enumerate(linha):
                        if mes_escolhido in cel:
                            col_inicio = j
                            break
                    break

            if linha_mes is None:
                return "❌ Mês não encontrado no arquivo."

            meses_lista = [m.lower() for m in meses]
            for j in range(col_inicio + 1, len(df_full.columns)):
                cel = str(df_full.iloc[linha_mes, j] or "").lower()
                if any(m in cel for m in meses_lista) or "total" in cel:
                    col_fim = j
                    break

            if col_fim is None:
                col_fim = len(df_full.columns)

            header_row    = linha_mes + 1
            df_base       = pd.read_excel(caminho, header=header_row)
            descricao_col = df_base.columns[1]
            df_valores    = df_base.iloc[:, col_inicio:col_fim]
            df            = pd.concat([df_base[[descricao_col]], df_valores], axis=1)
            df            = df.dropna(how="all")

            vendedores = df.columns[1:]
            arquivos   = []

            for vendedor in vendedores:
                dados          = {}
                total_planilha = 0

                for i in range(len(df)):
                    descricao   = str(df.iloc[i][descricao_col]).strip()
                    valor_bruto = df.iloc[i][vendedor]
                    valor       = tratar_valor(valor_bruto)

                    if descricao.upper() == "TOTAL":
                        if valor is not None:
                            total_planilha = valor
                        continue

                    if valor is not None and valor != 0:
                        dados[descricao] = valor

                if dados:
                    imagem_vendedor = encontrar_imagem(vendedor, mapa_imagens)
                    dados_bancarios = encontrar_dados_bancarios(vendedor, mapa_banco)

                    arquivo = gerar_recibo(
                        vendedor, dados, request.form["mes"],
                        total_planilha, data_recibo,
                        imagem_vendedor, dados_bancarios
                    )
                    arquivos.append(arquivo)

            zip_path = os.path.join(UPLOAD_FOLDER, "Recibos.zip")
            with zipfile.ZipFile(zip_path, "w") as zipf:
                for arq in arquivos:
                    zipf.write(arq, os.path.basename(arq))

            return send_file(zip_path, as_attachment=True)

        except Exception as e:
            import traceback
            erro = traceback.format_exc()
            print("❌ ERRO:", erro)
            return f"<pre>❌ Erro: {erro}</pre>"

    return render_template("index.html", meses=meses, mes_atual=mes_atual)


if __name__ == "__main__":
    app.run(debug=True)